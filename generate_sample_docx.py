import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_sample_docx():
    """Create a sample DOCX document with insurance policy information"""
    # Initialize Document object
    doc = Document()
    
    # Set document properties
    doc.core_properties.title = "National Parivar Mediclaim Plus Policy"
    doc.core_properties.author = "Insurance Company"
    
    # Add title
    title = doc.add_heading("National Parivar Mediclaim Plus Policy", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add policy sections
    add_section(doc, "1. DEFINITIONS", [
        "1.1 Hospital: An institution established for in-patient care and day care treatment of illness and/or injuries and which has been registered as a hospital with the local authorities under Clinical Establishments (Registration and Regulation) Act 2010 or under enactments specified under the Schedule of Section 56(1) and the said act Or complies with all minimum criteria as under:",
        "- has qualified nursing staff under its employment round the clock;",
        "- has at least 10 in-patient beds in towns having a population of less than 10,00,000 and at least 15 in-patient beds in all other places;",
        "- has qualified medical practitioner(s) in charge round the clock;",
        "- has a fully equipped operation theatre of its own where surgical procedures are carried out;",
        "- maintains daily records of patients and makes these accessible to the insurance company's authorized personnel."
    ])
    
    add_section(doc, "2. COVERAGE", [
        "2.1 Hospitalization: The Company shall indemnify medical expenses incurred for Hospitalization of the Insured Person during the Policy year, up to the Sum Insured and Cumulative Bonus specified in the policy schedule.",
        "2.2 AYUSH Treatment: The Company shall indemnify medical expenses incurred for inpatient care treatment under Ayurveda, Yoga and Naturopathy, Unani, Siddha and Homeopathy systems of medicines during the Policy Period up to the limit of sum insured as specified in the policy schedule in any AYUSH Hospital.",
        "2.3 Maternity Expenses: The Company shall indemnify maternity expenses incurred for the delivery of a child and/or expenses related to medically necessary and lawful termination of pregnancy during the policy period, subject to the following conditions:",
        "- The female insured person must have been continuously covered for at least 24 months.",
        "- This benefit is limited to two deliveries or terminations during the lifetime of the female insured person.",
        "2.4 Organ Donor Expenses: The Company shall indemnify medical expenses incurred by the Insured Person towards in-patient hospitalization of an organ donor for harvesting the organ, provided that:",
        "- The organ donor is any person whose organ has been made available in accordance with the Transplantation of Human Organs Act 1994 and the organ donated is for the use of the Insured Person.",
        "- The Company has accepted an inpatient Hospitalization claim for the insured member under the policy."
    ])
    
    add_section(doc, "3. WAITING PERIODS AND EXCLUSIONS", [
        "3.1 Pre-Existing Diseases: All Pre-Existing Diseases and their direct complications shall not be covered until 36 months of continuous coverage have elapsed since inception of the first policy with the Company.",
        "3.2 Specific Waiting Period: The expenses related to the treatment of any listed conditions, surgeries/treatments shall not be covered until the specified waiting period has elapsed since policy inception:",
        "- 24 months: Benign ENT disorders, tonsillectomy, adenoidectomy, mastoidectomy, tympanoplasty, hysterectomy, all internal and external benign tumours, cysts, polyps of any kind, benign breast disorders, benign prostatic hypertrophy, hernia of all types, hydrocele, fissures/fistula in anus, piles, pilonidal sinus, chronic sinusitis, chronic suppurative otitis media, deviated nasal septum, gout and rheumatism, gout and rheumatism, calculus diseases of gall bladder including cholecystitis, calculus diseases of urogenital system, all types of migraine/tension headache, SLE, varicose veins and varicose ulcers, congenital internal diseases.",
        "- 24 months: All treatments for joint replacement unless arising from accident, age-related osteoarthritis and osteoporosis.",
        "- 24 months: Cataract."
    ])
    
    add_section(doc, "4. GENERAL CONDITIONS", [
        "4.1 Grace Period: The Policy may be renewed by mutual consent. The Company shall not however be bound to give notice that it is due for renewal. A grace period of thirty days is allowed following the expiry of the policy period to maintain continuity of coverage. No coverage is available during the grace period. However, no coverage is available for expenses incurred during the break period.",
        "4.2 No Claim Discount: The insured shall be entitled for a No Claim Discount of 5% on the base premium, on renewal of a claim free policy, for a one-year policy term. The aggregate of such discount shall not exceed 5% of the total base premium.",
        "4.3 Preventive Health Check-up: Expenses incurred for preventive health check-up will be reimbursed at the end of a block of two continuous policy years, subject to a maximum of the amount specified in the Table of Benefits, provided no claims have been made during this period and the policy has been renewed without a break.",
        "4.4 Room Rent Capping: For Plan A, the daily room rent is capped at 1% of the Sum Insured, and ICU charges are capped at 2% of the Sum Insured. These limits do not apply if the treatment is for a listed procedure in a Preferred Provider Network (PPN)."
    ])
    
    # Save the document
    output_dir = "data"
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    output_path = os.path.join(output_dir, "sample_policy.docx")
    doc.save(output_path)
    
    print(f"Sample DOCX created at: {output_path}")
    return output_path

def add_section(doc, title, content):
    """Add a section to the document"""
    # Add section title
    heading = doc.add_heading(title, level=1)
    
    # Add content
    for line in content:
        if line.startswith("-"):
            # Add as bullet point
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(line[2:])  # Remove the dash and space
        else:
            # Add as normal paragraph
            p = doc.add_paragraph()
            p.add_run(line)
    
    # Add space after section
    doc.add_paragraph()

def main():
    """Main function to create sample documents"""
    print("=== Generating Sample Word Document ===")
    
    try:
        # Check if python-docx is installed
        import docx
    except ImportError:
        print("python-docx library not found. Installing...")
        import subprocess
        import sys
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
        print("python-docx installed successfully.")
    
    # Create sample DOCX
    docx_path = create_sample_docx()
    
    print("\nSample Word document generated successfully!")
    print(f"DOCX: {docx_path}")

if __name__ == "__main__":
    main()